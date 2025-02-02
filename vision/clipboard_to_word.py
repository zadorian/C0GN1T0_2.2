from PIL import ImageGrab, Image
import tempfile
import os
import base64
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from AI_models.ai_vision import VisionAnalyzer
from extract_patterns import ExtractionPattern

class WordFormatter:
    def __init__(self, output_file="extracted_content.docx"):
        self.doc = Document()
        self.output_file = output_file
        self._setup_document()
    
    def _setup_document(self):
        """Set up default document styles"""
        # Default font
        style = self.doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # Heading styles
        for level in range(1, 4):
            style = self.doc.styles[f'Heading {level}']
            style.font.name = 'Calibri'
            style.font.size = Pt(14 - level)  # Decreasing size for each level
            style.font.bold = True
    
    def add_heading(self, text, level=1):
        """Add a heading with specified level"""
        self.doc.add_heading(text, level=level)
    
    def add_paragraph(self, text, indent=0):
        """Add a paragraph with optional indentation"""
        p = self.doc.add_paragraph()
        if indent:
            p.paragraph_format.left_indent = Inches(indent * 0.25)
        p.add_run(text)
    
    def add_bullet_point(self, text, level=0):
        """Add a bullet point with specified level"""
        p = self.doc.add_paragraph(style='List Bullet')
        if level:
            p.paragraph_format.left_indent = Inches(level * 0.25)
        p.add_run(text)
    
    def add_numbered_item(self, text, level=0):
        """Add a numbered item with specified level"""
        p = self.doc.add_paragraph(style='List Number')
        if level:
            p.paragraph_format.left_indent = Inches(level * 0.25)
        p.add_run(text)
    
    def add_section_break(self):
        """Add a section break"""
        self.doc.add_paragraph()
    
    def save(self):
        """Save the document"""
        self.doc.save(self.output_file)
        print(f"\nDocument saved as: {self.output_file}")

def extract_and_format(pattern_name: str = None, output_file: str = None):
    """
    Extract text from clipboard image and format it into a Word document.
    
    Args:
        pattern_name: Optional name of extraction pattern to use
        output_file: Optional output file name
    """
    try:
        # Get image from clipboard
        clipboard_image = ImageGrab.grabclipboard()
        
        if not isinstance(clipboard_image, Image.Image):
            print("No image found in clipboard. Please copy an image first.")
            return None
            
        # Save to temporary file as JPEG
        with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as tmp_file:
            clipboard_image = clipboard_image.convert("RGB")
            clipboard_image.save(tmp_file.name, format='JPEG', quality=95)
            filename = tmp_file.name

        # Read and encode the image
        with open(filename, "rb") as image_file:
            image_data = base64.b64encode(image_file.read()).decode('utf-8')
        
        # Clean up the temporary file
        os.remove(filename)

        # Get extraction pattern
        if pattern_name:
            pattern_manager = ExtractionPattern()
            pattern = pattern_manager.get_pattern(pattern_name)
            if not pattern:
                print(f"Pattern '{pattern_name}' not found. Using default pattern.")
                pattern = get_default_pattern()
        else:
            pattern = get_default_pattern()
        
        # Initialize Word formatter
        word_formatter = WordFormatter(output_file or "extracted_content.docx")
        
        # Add pattern information to document
        word_formatter.add_heading("Extracted Content", level=1)
        word_formatter.add_paragraph(f"Pattern: {pattern_name or 'Default'}")
        word_formatter.add_section_break()
        
        # Create extraction prompt using pattern
        extraction_prompt = f"""Extract and structure ALL text visible in this image according to this pattern:

{pattern}

Additionally, format your response with these markers to indicate document structure:
- Use [H1] for main headings
- Use [H2] for subheadings
- Use [H3] for sub-subheadings
- Use [B] for bullet points
- Use [N] for numbered items
- Use [I1], [I2], etc. for different indentation levels
- Use [BREAK] for section breaks

Return ONLY the extracted and structured text with these markers."""

        # Initialize VisionAnalyzer with Claude
        analyzer = VisionAnalyzer()
        extracted_text = analyzer.analyze(
            image_data,
            extraction_prompt,
            model="claude"
        )
        
        if not extracted_text:
            print("No text was extracted.")
            return
        
        # Process the extracted text and format it in Word
        for line in extracted_text.split('\n'):
            line = line.strip()
            if not line:
                continue
                
            # Process formatting markers
            if line.startswith('[H1]'):
                word_formatter.add_heading(line[4:].strip(), 1)
            elif line.startswith('[H2]'):
                word_formatter.add_heading(line[4:].strip(), 2)
            elif line.startswith('[H3]'):
                word_formatter.add_heading(line[4:].strip(), 3)
            elif line.startswith('[B]'):
                word_formatter.add_bullet_point(line[3:].strip())
            elif line.startswith('[N]'):
                word_formatter.add_numbered_item(line[3:].strip())
            elif line.startswith('[I'):
                level = int(line[2]) - 1
                word_formatter.add_paragraph(line[4:].strip(), indent=level)
            elif line.startswith('[BREAK]'):
                word_formatter.add_section_break()
            else:
                word_formatter.add_paragraph(line)
        
        # Save the document
        word_formatter.save()
        
        return True

    except Exception as e:
        print(f"Error during extraction and formatting: {e}")
        return None

def get_default_pattern():
    """Returns the default text extraction pattern"""
    return """Extract and structure text following these rules:
1. Identify and mark main headings with [H1]
2. Mark subheadings with [H2] and sub-subheadings with [H3]
3. Format lists using:
   - [B] for bullet points
   - [N] for numbered items
4. Use indentation markers [I1], [I2], etc. for nested content
5. Insert [BREAK] between major sections
6. Preserve exact text and formatting:
   - CAPITALIZATION
   - *emphasis*
   - Numerical formats
   - Special characters"""

if __name__ == "__main__":
    print("Clipboard to Word Extractor")
    print("Commands:")
    print("pattern list - Show available patterns")
    print("pattern use [name] [output_file] - Use specific pattern")
    print("extract [output_file] - Extract with default pattern")
    print("quit - Exit program")
    
    while True:
        cmd = input("\nEnter command: ").strip().lower()
        
        if cmd == "quit":
            break
        elif cmd == "pattern list":
            pattern_manager = ExtractionPattern()
            pattern_manager.list_patterns()
        elif cmd.startswith("pattern use "):
            parts = cmd[11:].strip().split()
            pattern_name = parts[0]
            output_file = parts[1] if len(parts) > 1 else None
            print(f"Extracting text using pattern '{pattern_name}'...")
            extract_and_format(pattern_name, output_file)
        elif cmd.startswith("extract"):
            parts = cmd.split()
            output_file = parts[1] if len(parts) > 1 else None
            print("Extracting text with default pattern...")
            extract_and_format(output_file=output_file)
        else:
            print("Unknown command. Try: pattern list, pattern use [name] [output_file], extract [output_file], or quit") 