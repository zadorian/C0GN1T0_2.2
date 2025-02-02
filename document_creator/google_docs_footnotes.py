from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from docx import Document
import os
import pickle
import re
import tempfile
from typing import List, Tuple, Dict

class GoogleDocsFootnotes:
    def __init__(self, template_path: str = None):
        """
        Initialize Google Docs API with necessary credentials.
        
        Args:
            template_path: Path to the Word document template to use
        """
        self.SCOPES = ['https://www.googleapis.com/auth/documents', 'https://www.googleapis.com/auth/drive.file']
        self.creds = None
        self.template_path = template_path
        self._setup_credentials()
        self.service = build('docs', 'v1', credentials=self.creds)
        self.drive_service = build('drive', 'v3', credentials=self.creds)
    
    def _setup_credentials(self):
        """Set up Google Docs API credentials."""
        # Check for existing token
        if os.path.exists('token.pickle'):
            with open('token.pickle', 'rb') as token:
                self.creds = pickle.load(token)
        
        # If no valid credentials available, let user log in
        if not self.creds or not self.creds.valid:
            if self.creds and self.creds.expired and self.creds.refresh_token:
                self.creds.refresh(Request())
            else:
                if not os.path.exists('credentials.json'):
                    raise FileNotFoundError(
                        "credentials.json not found! Please follow these steps:\n"
                        "1. Go to https://console.cloud.google.com/\n"
                        "2. Create a new project\n"
                        "3. Enable Google Docs API and Google Drive API\n"
                        "4. Go to Credentials\n"
                        "5. Create OAuth 2.0 Client ID credentials\n"
                        "6. Download and save as 'credentials.json' in this directory"
                    )
                
                flow = InstalledAppFlow.from_client_secrets_file('credentials.json', self.SCOPES)
                self.creds = flow.run_local_server(port=0)
                
                # Save credentials for future use
                with open('token.pickle', 'wb') as token:
                    pickle.dump(self.creds, token)

    def process_text_with_urls(self, text: str) -> List[Tuple[str, str]]:
        """
        Process text containing URLs in square brackets and convert to content with footnotes.
        
        Args:
            text: Text containing URLs in square brackets [https://example.com]
            
        Returns:
            List of (text_content, footnote_text) tuples
        """
        # Split text into lines
        lines = text.split('\n')
        result = []
        footnote_number = 1
        
        for line in lines:
            if not line.strip():
                # Add empty lines without footnotes
                result.append((line, ""))
                continue
                
            # Find all URLs in square brackets
            urls = re.findall(r'\[(https?://[^\]]+)\]', line)
            if urls:
                # Replace each URL with a footnote number and store the URL
                modified_line = line
                for url in urls:
                    modified_line = modified_line.replace(f'[{url}]', f'[{footnote_number}]')
                    result.append((modified_line, url))
                    footnote_number += 1
            else:
                # No URLs in this line, add without footnote
                result.append((line, ""))
        
        return result

    def create_doc_from_template(self, title: str, content_sections: Dict[str, str]) -> str:
        """
        Create a Google Doc from a Word template and populate it with content.
        
        Args:
            title: The title of the document
            content_sections: Dictionary mapping section names to their content
                            e.g., {'ASSIGNMENT': 'assignment text', 'BACKGROUND': 'background text'}
        
        Returns:
            URL of the created document
        """
        if not self.template_path:
            raise ValueError("No template path provided")

        # First create a temporary copy of the template and modify it
        temp_doc = Document(self.template_path)
        
        # Process each section and replace content
        for section_name, content in content_sections.items():
            # Process URLs in the content
            content_with_footnotes = self.process_text_with_urls(content)
            
            # Find the section in the template
            for paragraph in temp_doc.paragraphs:
                if section_name in paragraph.text:
                    # Clear existing content after the section header
                    start_idx = list(temp_doc.paragraphs).index(paragraph) + 1
                    
                    # Add the new content
                    for text, footnote in content_with_footnotes:
                        p = temp_doc.add_paragraph(text)
                        if footnote:
                            p.add_footnote(footnote)

        # Save the modified template to a temporary file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
            temp_doc.save(tmp_file.name)
            temp_path = tmp_file.name

        try:
            # Upload the modified template to Google Drive
            file_metadata = {'name': title, 'mimeType': 'application/vnd.google-apps.document'}
            media = MediaFileUpload(
                temp_path,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                resumable=True
            )
            
            file = self.drive_service.files().create(
                body=file_metadata,
                media_body=media,
                fields='id'
            ).execute()
            
            document_id = file.get('id')
            return f"https://docs.google.com/document/d/{document_id}/edit"
            
        finally:
            # Clean up temporary file
            os.unlink(temp_path)

def main():
    # Use the correct template path
    template_path = "template/Subcontractor Report Template - final 1.docx"
    
    # Verify template exists
    if not os.path.exists(template_path):
        print(f"Error: Template file not found at {template_path}")
        print("Please make sure the template file exists in the template directory")
        return
        
    try:
        gdocs = GoogleDocsFootnotes(template_path)
        
        # Example content for different sections
        content_sections = {
            "A. ASSIGNMENT": """I was given an assignment to conduct checks on the following subjects:

**Date assigned:**
**Date due:** 27 January 2025
**Country:** Somalia""",
            
            "B. BACKGROUND INFORMATION": """#### Company Background Details

**Name:** Talosan Engineering Consultant
**Short Name:** Talosan Engineering
**Business Addresses:** Jicir Tower Near Total, Hargeisa, Somalia [https://www.facebook.com/Taloasanconsultingfirm/]
**Registration Number:** MPWR/REG/181/2017 [https://mpwr.gov.so/en/cm-business/talosan-engineering-consultant/]
**Category A License:** MPWR/REG/1060/2023 [https://mpwr.gov.so/members/talosan-engineering-consultant-2/]
**Company Type:** Limited liability company"""
        }
        
        print(f"\nCreating document from template: {template_path}")
        doc_url = gdocs.create_doc_from_template("Subcontractor Report", content_sections)
        print(f"\nDocument created successfully!")
        print(f"You can view it at: {doc_url}")
        
    except Exception as e:
        print(f"\nError creating document: {str(e)}")
        print("Please make sure:")
        print("1. The template file is not corrupted")
        print("2. You have proper Google API credentials set up")
        print("3. You have internet connectivity")

if __name__ == "__main__":
    main() 