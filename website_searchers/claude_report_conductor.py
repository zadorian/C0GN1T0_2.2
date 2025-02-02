import os
from anthropic import Anthropic
from typing import Dict, List, Optional, Any
import json
from datetime import datetime
from pathlib import Path
import logging
from docx import Document
from dotenv import load_dotenv
from PIL import ImageGrab, Image
import tempfile
import base64
import time
from pynput import keyboard
import shutil
import io
import aiohttp
import asyncio
import sys

# Add parent directory to Python path
sys.path.append(str(Path(__file__).parent.parent))
from config import config
from claude_pdf_ocr import ClaudeProcessor
from AI_models.ai_vision import VisionAnalyzer
from clipboard_extract import extract_text_from_clipboard
from search_engines.exa_exact import exa_search

# Add the company_search directory to Python path
sys.path.append(str(Path(__file__).parent.parent / "company_search"))
from aleph_search2 import AlephAPI

# Remove these imports as we're not using Azure anymore
# from aleph_doc_ocr import AlephPDFProcessor  # Remove this
# from azure_doc import AzureOCRProcessor     # Remove this

# Load environment variables from .env file
load_dotenv()

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class AlephSearcher:
    def __init__(self):
        self.api_key = "1c0971afa4804c2aafabb125c79b275e"
        self.base_url = "https://aleph.occrp.org/api/2/"
        self.headers = {
            'Authorization': f'ApiKey {self.api_key}',
            'Accept': 'application/json'
        }

    async def search_entity(self, query: str) -> Dict:
        """Search for an entity and its network in Aleph"""
        try:
            async with aiohttp.ClientSession() as session:
                # Initial entity search
                search_url = f"{self.base_url}entities"
                params = {
                    "q": query,
                    "filter:schema": "Company",
                    "limit": 10
                }
                
                async with session.get(search_url, headers=self.headers, params=params) as response:
                    response.raise_for_status()
                    results = await response.json()
                    
                    if not results.get('results'):
                        # Try without company filter
                        params.pop('filter:schema')
                        async with session.get(search_url, headers=self.headers, params=params) as broad_response:
                            broad_response.raise_for_status()
                            results = await broad_response.json()
                    
                    if not results.get('results'):
                        return {}
                        
                    entity = results['results'][0]
                    entity_id = entity['id']
                    
                    # Get full entity details
                    entity_url = f"{self.base_url}entities/{entity_id}"
                    params = {
                        'include': 'properties,schema,names,addresses,countries,identifiers,tags',
                        'expand': 'true'
                    }
                    
                    async with session.get(entity_url, headers=self.headers, params=params) as entity_response:
                        entity_response.raise_for_status()
                        full_entity = await entity_response.json()
                        
                    # Get network information
                    network = await self._get_entity_network(session, entity_id)
                    network['entity'] = full_entity
                    
                    return network
                    
        except Exception as e:
            logger.error(f"Aleph search failed: {str(e)}")
            return {}

    async def _get_entity_network(self, session: aiohttp.ClientSession, entity_id: str) -> Dict:
        """Get the complete network for an entity"""
        network = {
            'directors': [],
            'owners': [],
            'related_entities': [],
            'documents': [],
            'statements': []
        }
        
        try:
            # Get directors
            dir_url = f"{self.base_url}entities"
            dir_params = {
                'filter:schema': 'Directorship',
                'filter:properties.organization': entity_id,
                'include': 'properties'
            }
            
            async with session.get(dir_url, headers=self.headers, params=dir_params) as dir_response:
                dir_response.raise_for_status()
                directorships = await dir_response.json()
                
                for d in directorships.get('results', []):
                    props = d.get('properties', {})
                    director_id = props.get('director', {}).get('id')
                    if director_id:
                        async with session.get(f"{self.base_url}entities/{director_id}", headers=self.headers) as d_response:
                            if d_response.status == 200:
                                director = await d_response.json()
                                director['position'] = props.get('role', 'Director')
                                network['directors'].append(director)
            
            # Get owners
            own_params = {
                'filter:schema': 'Ownership',
                'filter:properties.asset': entity_id,
                'include': 'properties'
            }
            
            async with session.get(dir_url, headers=self.headers, params=own_params) as own_response:
                own_response.raise_for_status()
                ownerships = await own_response.json()
                
                for o in ownerships.get('results', []):
                    props = o.get('properties', {})
                    owner_id = props.get('owner', {}).get('id')
                    if owner_id:
                        async with session.get(f"{self.base_url}entities/{owner_id}", headers=self.headers) as o_response:
                            if o_response.status == 200:
                                owner = await o_response.json()
                                owner['ownership_percentage'] = props.get('shareValue')
                                network['owners'].append(owner)
            
            # Get related documents
            doc_params = {
                'filter:schema': 'Pages',
                'filter:properties.mentions': entity_id
            }
            
            async with session.get(dir_url, headers=self.headers, params=doc_params) as doc_response:
                doc_response.raise_for_status()
                documents = await doc_response.json()
                network['documents'].extend(documents.get('results', []))
            
            return network
            
        except Exception as e:
            logger.error(f"Error getting entity network: {str(e)}")
            return network

class ExactStyleReportGenerator:
    def __init__(self):
        """Initialize report generator with template and output directory"""
        template_dir = Path(__file__).parent.parent / "template"
        reports_dir = Path(__file__).parent.parent / "reports"
        
        # Create directories if they don't exist
        template_dir.mkdir(parents=True, exist_ok=True)
        reports_dir.mkdir(parents=True, exist_ok=True)
        
        self.template_path = template_dir / "Subcontractor Report Template - final 1.docx"
        self.output_dir = reports_dir
        self.current_report = None
        self.current_report_path = None
        
        # Create template if it doesn't exist
        if not self.template_path.exists():
            # Create a basic template
            doc = Document()
            doc.add_heading('Subcontractor Report Template', 0)
            for section in ["Company Background", "Negative Media Findings", "Additional Documents", "Additional Information"]:
                doc.add_heading(section, level=1)
                doc.add_paragraph()
            doc.save(self.template_path)
            print(f"Created new template at {self.template_path}")

    def create_new_report(self, company_name: str) -> str:
        """Create a new report from template"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_name = "".join(c if c.isalnum() else "_" for c in company_name)
        report_name = f"{safe_name}_{timestamp}.docx"
        self.current_report_path = self.output_dir / report_name
        
        # Create new document with required sections
        doc = Document()
        doc.add_heading(f'Report: {company_name}', 0)
        
        # Always create these sections
        sections = [
            "Company Background",
            "Negative Media Findings",
            "Additional Documents",
            "Additional Information"
        ]
        
        for section in sections:
            heading = doc.add_heading(section, level=1)
            # Add a paragraph after each heading to ensure sections are distinct
            doc.add_paragraph()
        
        # Save the new report
        doc.save(self.current_report_path)
        print(f"\nCreated new report at: {self.current_report_path}")
        print(f"Report directory: {self.output_dir}")
        
        self.current_report = doc
        return str(self.current_report_path)

    def add_content_to_section(self, content: str, section: str):
        """Add content to a specific section in the report"""
        if not self.current_report or not self.current_report_path:
            raise ValueError("No report is currently open")

        # Find the section heading
        section_found = False
        for paragraph in self.current_report.paragraphs:
            if paragraph.text.strip() == section:
                # Add content after the section heading
                new_paragraph = self.current_report.add_paragraph()
                new_paragraph.text = content
                section_found = True
                break
        
        if not section_found:
            # Add section if it doesn't exist
            self.current_report.add_heading(section, level=1)
            new_paragraph = self.current_report.add_paragraph()
            new_paragraph.text = content
        
        # Save changes
        try:
            self.current_report.save(self.current_report_path)
            print(f"Saved content to section '{section}' in {self.current_report_path.name}")
        except Exception as e:
            print(f"Error saving report: {e}")
            raise

class ClipboardMonitor:
    def __init__(self, base_dir: str = "ingest"):
        self.base_dir = Path(base_dir)
        self.running = True
        self.setup_directories()
        self.last_clipboard_content = None
        self.processed_files = set()
        
    def setup_directories(self):
        """Create necessary directories for ingestion"""
        for dir_name in ["stomach", "uncertain"]:
            (self.base_dir / dir_name).mkdir(parents=True, exist_ok=True)
            
    def check_clipboard(self) -> Optional[Dict]:
        """Check clipboard for new content"""
        try:
            image = ImageGrab.grabclipboard()
            
            if isinstance(image, Image.Image):
                # Convert to bytes for comparison
                with io.BytesIO() as bio:
                    image.save(bio, format='PNG')
                    current_content = bio.getvalue()
                
                if current_content != self.last_clipboard_content:
                    self.last_clipboard_content = current_content
                    
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    image_path = self.base_dir / f"clip_{timestamp}.png"
                    image.save(image_path)
                    
                    metadata = {
                        "timestamp": timestamp,
                        "path": str(image_path),
                        "processed": False,
                        "source_url": None,
                        "user_comment": "Auto-captured from clipboard"
                    }
                    return metadata
            return None
            
        except Exception as e:
            logger.error(f"Error checking clipboard: {e}")
            return None

    def check_ingest_folder(self) -> List[Dict]:
        """Check ingest folder for new files"""
        new_files = []
        try:
            for file_path in self.base_dir.glob("*"):
                if file_path.is_file() and str(file_path) not in self.processed_files:
                    metadata = {
                        "timestamp": datetime.now().strftime("%Y%m%d_%H%M%S"),
                        "path": str(file_path),
                        "processed": False,
                        "source_url": None,
                        "user_comment": f"Found in ingest folder: {file_path.name}"
                    }
                    new_files.append(metadata)
                    self.processed_files.add(str(file_path))
        except Exception as e:
            logger.error(f"Error checking ingest folder: {e}")
        return new_files

    def start_monitoring(self, callback):
        """Start continuous monitoring of both clipboard and folder"""
        print("\nAutomatic Monitor Active")
        print("Watching clipboard and ingest folder for new content")
        print("(Press Ctrl+C to stop)")
        
        try:
            while self.running:
                # Check clipboard
                clipboard_metadata = self.check_clipboard()
                if clipboard_metadata:
                    print("\nNew clipboard content detected!")
                    try:
                        asyncio.run(callback(clipboard_metadata))
                        print("Clipboard content processed successfully")
                    except Exception as e:
                        print(f"Error processing clipboard content: {e}")

                # Check folder
                new_files = self.check_ingest_folder()
                for file_metadata in new_files:
                    file_name = Path(file_metadata['path']).name
                    print(f"\nFound new file: {file_name}")
                    try:
                        print(f"Starting processing of {file_name}...")
                        asyncio.run(callback(file_metadata))
                        print(f"Successfully processed {file_name}")
                    except Exception as e:
                        print(f"Error processing {file_name}: {e}")

                time.sleep(1)  # Check every second
                
        except KeyboardInterrupt:
            print("\nStopping monitor...")
            self.running = False

class ClaudeConductor:
    def __init__(self):
        self.client = Anthropic(api_key=os.getenv('ANTHROPIC_API_KEY'))
        self.model = "claude-3-opus-20240229"
        self.report_generator = ExactStyleReportGenerator()
        self.clipboard_monitor = ClipboardMonitor()
        self.aleph_searcher = AlephAPI()
        self.claude_processor = ClaudeProcessor()
        self.vision_analyzer = VisionAnalyzer()
        self.collected_data = {
            "screenshots": [],
            "structured_data": {},
            "aleph_data": {},
            "exa_data": [],
            "ocr_data": []
        }

    async def search_subject(self, query: str):
        """Perform comprehensive search on a subject using Aleph, OCR, and Exa"""
        print(f"\nSearching for information about: {query}")
        
        # Create timestamp for this search
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Search Aleph
        print("\nSearching Aleph...")
        aleph_results = self.aleph_searcher.search_entity(query)
        if aleph_results:
            # Save Aleph results
            aleph_file = f"ingest/aleph_{timestamp}.json"
            with open(aleph_file, 'w', encoding='utf-8') as f:
                json.dump(aleph_results, f, ensure_ascii=False, indent=2)
            print(f"Saved Aleph results to {aleph_file}")
            
            # Process the file
            await asyncio.sleep(1)  # Give time for file to be written
            await self.process_screenshot({"path": aleph_file, "source": "aleph"})
        
        # Search Exa
        print("\nSearching Exa...")
        exa_search(query)  # This will save results to ingest folder
        await asyncio.sleep(1)  # Give time for file to be written
        
        # Process any Exa results files
        for file_path in Path("ingest").glob(f"{timestamp}_*.json"):
            if "exa" in file_path.name.lower():
                await self.process_screenshot({"path": str(file_path), "source": "exa"})
        
        print("\nSearch complete! All results have been saved to the ingest folder.")

    def _process_aleph_results(self, data: Dict) -> str:
        """Process Aleph search results"""
        analysis = []
        
        # Process entity information
        if entity := data.get('entity'):
            analysis.append(f"Entity: {entity.get('name')}")
            analysis.append(f"Type: {entity.get('schema')}")
            
            if props := entity.get('properties', {}):
                for key, value in props.items():
                    if value:
                        analysis.append(f"{key}: {value}")
        
        # Process directors
        if directors := data.get('directors'):
            analysis.append("\nDirectors:")
            for d in directors:
                analysis.append(f"- {d.get('name')} ({d.get('position', 'Director')})")
        
        # Process owners
        if owners := data.get('owners'):
            analysis.append("\nOwners:")
            for o in owners:
                share = f" ({o.get('ownership_percentage')}%)" if o.get('ownership_percentage') else ""
                analysis.append(f"- {o.get('name')}{share}")
        
        # Process documents
        if documents := data.get('documents'):
            analysis.append("\nRelated Documents:")
            for doc in documents:
                analysis.append(f"- {doc.get('title')}")
                if url := doc.get('links', {}).get('file'):
                    analysis.append(f"  URL: {url}")
        
        return "\n".join(analysis)

    def _process_exa_results(self, data: List[Dict]) -> None:
        """Process Exa search results and add them to the report"""
        if not self.report_generator.current_report:
            print("No active report - creating new one...")
            self.report_generator.create_new_report("Exa Search Results")

        # Find the "Negative Media Findings" section
        target_section = None
        for paragraph in self.report_generator.current_report.paragraphs:
            if paragraph.text.strip() == "Negative Media Findings":
                target_section = paragraph
                break

        if not target_section:
            print("Error: Could not find 'Negative Media Findings' section in report")
            return

        # Add search results after the section heading
        for result in data:
            # Add title and URL
            title_para = self.report_generator.current_report.add_paragraph()
            title_para.add_run(f"\n{result['title']}").bold = True
            url_para = self.report_generator.current_report.add_paragraph()
            url_para.add_run(f"Source: {result['url']}")
            
            # Add published date if available
            if result.get('published_date'):
                date_para = self.report_generator.current_report.add_paragraph()
                date_para.add_run(f"Published: {result['published_date']}")

            # Add highlights if available
            if result.get('highlights'):
                highlight_para = self.report_generator.current_report.add_paragraph()
                highlight_para.add_run("Key Highlights:").bold = True
                for highlight in result['highlights']:
                    self.report_generator.current_report.add_paragraph(highlight, style='List Bullet')

            # Add main text
            if result.get('text'):
                text_para = self.report_generator.current_report.add_paragraph()
                text_para.add_run("\nFull Text:").bold = True
                self.report_generator.current_report.add_paragraph(result['text'])

            # Add separator
            self.report_generator.current_report.add_paragraph("\n" + "-"*50 + "\n")

        # Save the updated report
        self.report_generator.current_report.save(self.report_generator.current_report_path)
        print(f"\nUpdated report saved to: {self.report_generator.current_report_path}")

    def _is_relevant_content(self, analysis: str) -> bool:
        """Determine if the analyzed content is relevant"""
        # Add your relevance criteria here
        return bool(analysis.strip())

    def _resize_image_if_needed(self, image_path: str, max_size_mb: int = 4) -> str:
        """Resize image if it's larger than max_size_mb"""
        try:
            # Get file size in MB
            file_size = Path(image_path).stat().st_size / (1024 * 1024)
            
            if file_size > max_size_mb:
                print(f"Image size: {file_size:.1f}MB - Resizing to fit {max_size_mb}MB limit...")
                
                # Open and resize image while maintaining aspect ratio
                with Image.open(image_path) as img:
                    # Calculate scaling factor
                    scale = (max_size_mb / file_size) ** 0.5
                    new_size = tuple(int(dim * scale) for dim in img.size)
                    
                    # Resize image
                    resized = img.resize(new_size, Image.Resampling.LANCZOS)
                    
                    # Save to temporary file
                    temp_path = f"{image_path}_resized.png"
                    resized.save(temp_path, "PNG", optimize=True, quality=85)
                    
                    print(f"Resized image saved: {Path(temp_path).name}")
                    return temp_path
            
            return image_path
            
        except Exception as e:
            logger.error(f"Error resizing image: {e}")
            return image_path

    async def process_screenshot(self, metadata: Dict):
        """Process a captured screenshot or file"""
        try:
            file_path = Path(metadata["path"])
            print(f"\n{'='*50}")
            print(f"NOW PROCESSING: {file_path.name}")
            print(f"Full path: {file_path}")
            print(f"Type: {metadata.get('source_type', 'Unknown')}")
            print(f"{'='*50}\n")
            
            # Keep track of original file path
            original_file_path = file_path
            
            analysis = None
            temp_files = []  # Track temporary files for cleanup
            
            try:
                # Handle JSON files from Aleph or Exa
                if file_path.suffix.lower() == '.json':
                    print("Analyzing JSON data...")
                    with open(file_path, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                    
                    # Determine source and process accordingly
                    if metadata.get("source") == "aleph":
                        print("Processing Aleph search results...")
                        analysis = self._process_aleph_results(data)
                        metadata["source_type"] = "aleph_search"
                    elif metadata.get("source") == "exa":
                        print("Processing Exa search results...")
                        self._process_exa_results(data)
                        metadata["source_type"] = "exa_search"
                    else:
                        print("Determining data source type...")
                        if isinstance(data, dict) and 'entity' in data:
                            analysis = self._process_aleph_results(data)
                            metadata["source_type"] = "aleph_search"
                        elif isinstance(data, list) and all(isinstance(r, dict) and 'url' in r and 'text' in r for r in data):
                            self._process_exa_results(data)
                            metadata["source_type"] = "exa_search"
                        else:
                            raise ValueError("Unknown JSON format")
                    
                    metadata["analysis"] = analysis
                    metadata["processed"] = True
                
                # Handle all other files (PDFs and images) with Vision
                else:
                    print("Processing with Vision Analyzer...")
                    
                    # Compress image if needed
                    if file_path.suffix.lower() in ['.png', '.jpg', '.jpeg']:
                        print("Compressing image before analysis...")
                        with Image.open(file_path) as img:
                            # Convert to RGB and resize if too large
                            img = img.convert("RGB")
                            
                            # Calculate target size to maintain aspect ratio
                            max_dimension = 2000  # Max width/height
                            ratio = min(max_dimension / img.width, max_dimension / img.height)
                            if ratio < 1:  # Only resize if image is too large
                                new_size = (int(img.width * ratio), int(img.height * ratio))
                                img = img.resize(new_size, Image.Resampling.LANCZOS)
                                print(f"Resized image to {new_size}")
                            
                            # Save as PNG to maintain quality
                            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
                                img.save(tmp_file.name, format='PNG', optimize=True)
                                file_path = Path(tmp_file.name)
                                temp_files.append(file_path)
                                print(f"Processed image saved as PNG")
                    
                    # Use VisionAnalyzer for both PDFs and images
                    prompt = "Extract all relevant business information from this document. Include any URLs, dates, names, and key details visible."
                    try:
                        analysis = self.vision_analyzer.analyze(
                            str(file_path),
                            prompt,
                            model="claude"
                        )
                        
                        # Check if analysis contains error message
                        if "Error analyzing image with Claude Vision" in analysis:
                            raise Exception(f"Vision API Error: {analysis}")
                            
                        print("Analysis complete")
                        metadata["analysis"] = analysis
                        metadata["source_type"] = "pdf" if file_path.suffix.lower() == '.pdf' else "image"
                        metadata["processed"] = True

                        # Only proceed with report addition if we have valid analysis
                        if analysis and not analysis.startswith("Error"):
                            print("\nANALYZED TEXT:")
                            print("─" * 50)
                            print(analysis)
                            print("─" * 50)
                            print("\nAdding to report...")
                            
                            # Create a new report if none exists
                            if not self.report_generator.current_report:
                                company_name = "Report_" + datetime.now().strftime("%Y%m%d_%H%M%S")
                                self.report_generator.create_new_report(company_name)
                            
                            # Add to report with appropriate styling
                            await self._add_to_report(analysis, metadata["source_type"])
                            print("Added to report successfully!")

                            # Move original file to stomach folder only if processing succeeded
                            if original_file_path.exists():
                                stomach_path = self.clipboard_monitor.base_dir / "stomach" / original_file_path.name
                                print(f"Moving original file to stomach: {stomach_path.name}")
                                shutil.move(str(original_file_path), str(stomach_path))
                                metadata["path"] = str(stomach_path)
                            
                            # Add to collected data if relevant
                            if self._is_relevant_content(metadata["analysis"]):
                                self.collected_data["screenshots"].append(metadata)
                                print("Added to collected data")
                            else:
                                print("Content deemed not relevant")
                        else:
                            raise Exception("No valid analysis produced")
                            
                    except Exception as vision_error:
                        print(f"Vision analysis failed: {vision_error}")
                        # Move file to uncertain since processing failed
                        if original_file_path.exists():
                            uncertain_path = self.clipboard_monitor.base_dir / "uncertain" / original_file_path.name
                            print(f"Moving file to uncertain due to failed analysis: {uncertain_path.name}")
                            shutil.move(str(original_file_path), str(uncertain_path))
                            metadata["path"] = str(uncertain_path)
                        raise  # Re-raise to trigger outer exception handler

            finally:
                # Clean up any temporary files
                for temp_file in temp_files:
                    try:
                        if temp_file.exists():
                            os.remove(temp_file)
                    except Exception as e:
                        print(f"Warning: Could not remove temporary file {temp_file}: {e}")
            
        except Exception as e:
            logger.error(f"Error processing file: {e}")
            try:
                if original_file_path.exists():
                    uncertain_path = self.clipboard_monitor.base_dir / "uncertain" / original_file_path.name
                    print(f"Error occurred - moving to uncertain: {uncertain_path.name}")
                    shutil.move(str(original_file_path), str(uncertain_path))
                    metadata["path"] = str(uncertain_path)
            except Exception as move_error:
                logger.error(f"Error moving file to uncertain: {move_error}")

    async def _add_to_report(self, content: str, source_type: str):
        """Add content to the current report document with appropriate styling"""
        try:
            # Determine section based on source_type
            section = ""
            if source_type == "aleph_search":
                section = "Company Background"
            elif source_type == "exa_search":
                section = "Negative Media Findings"
            elif source_type == "pdf":
                section = "Additional Documents"
            else:
                section = "Additional Information"

            # Add content to Word doc
            self.report_generator.add_content_to_section(content, section)
            print(f"Added content to report section: {section}")
            
        except Exception as e:
            logger.error(f"Error adding to report: {e}")
            raise

    def start_collection(self):
        """Start monitoring both clipboard and ingest folder"""
        print("\nStarting automatic monitoring...")
        self.clipboard_monitor.start_monitoring(self.process_screenshot)

    def process_raw_input(self, raw_data: Dict) -> Dict:
        """
        Have Claude process and structure raw input data.
        """
        try:
            response = self.client.messages.create(
                model=self.model,
                max_tokens=4096,
                temperature=0,
                system="You are a report data processor. Your task is to:\n"
                       "1. Analyze raw input data about companies/individuals\n"
                       "2. Structure it according to the exact Subcontractor Report template sections\n"
                       "3. Extract and verify source information for footnotes\n"
                       "4. Flag any missing required information\n"
                       "5. Format data appropriately for each section",
                messages=[{
                    "role": "user",
                    "content": f"""Process this raw data into structured sections for a Subcontractor Report:
                    {json.dumps(raw_data, indent=2)}
                    
                    Return a JSON object with these sections:
                    1. header (report date, subject name)
                    2. assignment_details
                    3. company_background (if applicable)
                    4. individual_background (if applicable)
                    5. negative_media_findings
                    6. source_footnotes (mapping each piece of information to its source)
                    
                    Return ONLY the JSON object, no additional text or explanation.
                    """
                }]
            )
            
            # Extract JSON from the response content
            content = response.content[0].text
            # Find the JSON object in the content
            start_idx = content.find('{')
            end_idx = content.rfind('}') + 1
            if start_idx == -1 or end_idx == 0:
                raise ValueError("No JSON object found in response")
            json_str = content[start_idx:end_idx]
            structured_data = json.loads(json_str)
            return self._validate_structured_data(structured_data)
        except Exception as e:
            logger.error(f"Error processing data with Claude: {e}")
            logger.error(f"Response content: {response.content if 'response' in locals() else 'No response'}")
            raise

    def process_source_documents(self, documents: List[Dict[str, str]]) -> Dict:
        """
        Have Claude extract relevant information from source documents.
        """
        try:
            documents_content = "\n\n".join([
                f"Document: {doc.get('title', 'Untitled')}\nContent:\n{doc['content']}"
                for doc in documents
            ])
            response = self.client.messages.create(
                model=self.model,
                max_tokens=4096,
                temperature=0,
                system="You are a document analyzer. Extract and structure information for reports.",
                messages=[{
                    "role": "user",
                    "content": f"""Extract relevant information from these documents for a Subcontractor Report:
                    {documents_content}
                    
                    Focus on:
                    1. Company/individual details
                    2. Negative media findings
                    3. Source information for footnotes
                    4. Key dates and events
                    
                    Return ONLY a JSON object with the structured information, no additional text.
                    """
                }]
            )
            
            # Extract JSON from the response content
            content = response.content[0].text
            # Find the JSON object in the content
            start_idx = content.find('{')
            end_idx = content.rfind('}') + 1
            if start_idx == -1 or end_idx == 0:
                raise ValueError("No JSON object found in response")
            json_str = content[start_idx:end_idx]
            return json.loads(json_str)
        except Exception as e:
            logger.error(f"Error processing documents with Claude: {e}")
            logger.error(f"Response content: {response.content if 'response' in locals() else 'No response'}")
            raise

    def validate_report_content(self, report_data: Dict) -> List[str]:
        """
        Have Claude validate report content and identify issues.
        """
        try:
            response = self.client.messages.create(
                model=self.model,
                max_tokens=4096,
                temperature=0,
                system="You are a report validator. Check for completeness and accuracy.",
                messages=[{
                    "role": "user",
                    "content": f"""Validate this report data against Subcontractor Report requirements:
                    {json.dumps(report_data, indent=2)}
                    
                    Check for:
                    1. Missing required fields
                    2. Inconsistent information
                    3. Proper source citations
                    4. Format compliance
                    
                    Return ONLY a JSON array of validation issues, no additional text.
                    """
                }]
            )
            
            # Extract JSON from the response content
            content = response.content[0].text
            # Find the JSON array in the content
            start_idx = content.find('[')
            end_idx = content.rfind(']') + 1
            if start_idx == -1 or end_idx == 0:
                raise ValueError("No JSON array found in response")
            json_str = content[start_idx:end_idx]
            return json.loads(json_str)
        except Exception as e:
            logger.error(f"Error validating report with Claude: {e}")
            raise

    def _validate_structured_data(self, data: Dict) -> Dict:
        """
        Validate the structure of processed data.
        """
        required_sections = ['header', 'assignment_details', 'source_footnotes']
        missing_sections = [section for section in required_sections if section not in data]
        if missing_sections:
            raise ValueError(f"Missing required sections: {missing_sections}")
        return data

    async def generate_report(self, 
                       input_data: Optional[Dict] = None,
                       source_documents: Optional[List[Dict]] = None) -> str:
        """Generate report from collected data"""
        try:
            # Create new report from template
            company_name = input_data.get("header", {}).get("subject_name", "Unknown")
            report_path = self.report_generator.create_new_report(company_name)
            
            # Process and add all content
            if self.collected_data["screenshots"]:
                for screenshot in self.collected_data["screenshots"]:
                    await self._add_to_report(
                        screenshot["analysis"],
                        screenshot["source_type"]
                    )
            
            if input_data:
                for section, content in input_data.items():
                    if section not in ["header", "assignment_details"]:
                        await self._add_to_report(content, section)
            
            return report_path
            
        except Exception as e:
            logger.error(f"Error generating report: {e}")
            raise

def main():
    conductor = ClaudeConductor()
    
    # Add command to show collected data
    def print_collected_data():
        print("\nCollected Data Summary:")
        print("-" * 30)
        for category, items in conductor.collected_data.items():
            print(f"\n{category.upper()}:")
            if isinstance(items, list):
                for item in items:
                    print(f"- {item.get('source_type', 'unknown')}: {item.get('path')}")
            elif isinstance(items, dict):
                for key, value in items.items():
                    print(f"- {key}: {value}")
    
    # Use the correct pynput keyboard listener
    def on_press(key):
        try:
            if key.char == 'p':
                print_collected_data()
        except AttributeError:
            pass
    
    # Start keyboard listener
    listener = keyboard.Listener(on_press=on_press)
    listener.start()
    
    print("\nStarting automatic monitoring...")
    print("Watching clipboard and ingest folder for new content")
    print("(Press Ctrl+C to stop)")
    print("(Press P to show collected data)")
    
    # Start monitoring immediately
    conductor.start_collection()

if __name__ == "__main__":
    main()