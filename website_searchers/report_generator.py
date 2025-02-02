import os
from typing import Dict, List, Optional, Literal
import json
from datetime import datetime
from pathlib import Path
import logging
from docx import Document
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
import pickle
import time

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class ExactStyleReportGenerator:
    def __init__(self):
        """Initialize with Google Docs support."""
        self.google_creds = None
        self._setup_google_docs()

    def _setup_google_docs(self):
        """Set up Google Docs API credentials."""
        SCOPES = ['https://www.googleapis.com/auth/documents', 
                 'https://www.googleapis.com/auth/drive.file']
        creds = None
        
        # Check for credentials.json
        if not os.path.exists('credentials.json'):
            raise FileNotFoundError(
                "credentials.json not found! Please follow these steps:\n"
                "1. Go to https://console.cloud.google.com/\n"
                "2. Create a new project\n"
                "3. Enable Google Docs API and Google Drive API\n"
                "4. Go to Credentials\n"
                "5. Create OAuth 2.0 Client ID credentials\n"
                "6. Download and save as 'credentials.json' in this directory"
            )
        
        # Check for existing token
        if os.path.exists('token.pickle'):
            try:
                with open('token.pickle', 'rb') as token:
                    creds = pickle.load(token)
                logger.info("Loaded existing Google credentials from token.pickle")
            except Exception as e:
                logger.warning(f"Error loading token.pickle: {e}")
                creds = None
        
        # If no valid credentials available, let the user log in
        if not creds or not creds.valid:
            if creds and creds.expired and creds.refresh_token:
                try:
                    creds.refresh(Request())
                    logger.info("Refreshed expired Google credentials")
                except Exception as e:
                    logger.warning(f"Error refreshing credentials: {e}")
                    creds = None
            
            if not creds:
                try:
                    print("\nOpening browser for Google authentication...")
                    print("Please log in with your Google account and grant access.")
                    flow = InstalledAppFlow.from_client_secrets_file(
                        'credentials.json', SCOPES)
                    creds = flow.run_local_server(port=0)
                    logger.info("Successfully obtained new Google credentials")
                    
                    # Save the credentials for the next run
                    with open('token.pickle', 'wb') as token:
                        pickle.dump(creds, token)
                    logger.info("Saved new credentials to token.pickle")
                except Exception as e:
                    raise RuntimeError(
                        f"Error during Google authentication: {e}\n"
                        "Make sure you have a valid credentials.json file and "
                        "are able to access a web browser for authentication."
                    )
        
        self.google_creds = creds

    def generate_report(self, structured_data: Dict) -> Dict[str, str]:
        """Generate report through the complete workflow."""
        # First generate Word doc with exact styling
        initial_word_path = self._generate_word_report(structured_data)
        logger.info(f"Initial Word document created: {initial_word_path}")
        
        # Convert to Google Docs and add footnotes
        gdoc_url = self._convert_to_google_docs(initial_word_path, structured_data)
        logger.info(f"Google Doc created with footnotes: {gdoc_url}")
        
        # Download back to Word with footnotes intact
        final_word_path = self._download_from_google_docs(gdoc_url)
        logger.info(f"Final Word document downloaded: {final_word_path}")
        
        return {
            'initial_word_path': initial_word_path,
            'gdoc_url': gdoc_url,
            'final_word_path': final_word_path
        }

    def _generate_word_report(self, structured_data: Dict) -> str:
        """Generate a DOCX report with exact styling."""
        # Get template path relative to script location
        template_path = Path(__file__).parent.parent / "template" / "Subcontractor Report Template - final 1.docx"
        doc = Document(template_path)

        # Header
        header_data = structured_data.get("header", {})
        doc.add_heading(f"Subcontractor Report: {header_data.get('subject_name', '')}", level=1)
        doc.add_paragraph(f"Report Date: {header_data.get('report_date', '')}")

        # Assignment Details
        assignment_details = structured_data.get("assignment_details", {})
        doc.add_heading("Assignment Details", level=2)
        for key, value in assignment_details.items():
            doc.add_paragraph(f"{key.capitalize()}: {value}")

        # Company Background
        if "company_background" in structured_data:
            doc.add_heading("Company Background", level=2)
            doc.add_paragraph(structured_data["company_background"])

        # Individual Background
        if "individual_background" in structured_data:
            doc.add_heading("Individual Background", level=2)
            doc.add_paragraph(structured_data["individual_background"])

        # Negative Media Findings
        if "negative_media_findings" in structured_data:
            doc.add_heading("Negative Media Findings", level=2)
            for finding in structured_data["negative_media_findings"]:
                doc.add_paragraph(finding)

        # Source Footnotes (temporary section, will be properly added in Google Docs)
        source_footnotes = structured_data.get("source_footnotes", {})
        if source_footnotes:
            doc.add_heading("Source References", level=2)
            for key, source_info in source_footnotes.items():
                text = f"{key}: {source_info['url']} (accessed {source_info['access_date']})"
                doc.add_paragraph(text)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = f"Subcontractor_Report_{timestamp}.docx"
        doc.save(output_path)
        return output_path

    def _convert_to_google_docs(self, word_path: str, structured_data: Dict) -> str:
        """Convert Word doc to Google Docs and add footnotes intelligently."""
        service = build('docs', 'v1', credentials=self.google_creds)
        drive_service = build('drive', 'v3', credentials=self.google_creds)

        # Create a new Google Doc
        doc_metadata = {
            'title': f'Subcontractor Report - {datetime.now().strftime("%Y%m%d_%H%M%S")}',
            'mimeType': 'application/vnd.google-apps.document'
        }
        doc = service.documents().create(body=doc_metadata).execute()
        document_id = doc.get('documentId')

        # Upload Word content
        media = MediaFileUpload(word_path, 
                              mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                              resumable=True)
        
        file = drive_service.files().create(
            body=doc_metadata,
            media_body=media,
            fields='id'
        ).execute()

        # Wait for conversion
        time.sleep(2)

        # Process content and add footnotes
        requests = []
        current_index = 1  # Start after title

        # Get the footnotes from structured data
        footnotes = structured_data.get('report', {}).get('footnotes', [])
        
        # Process each section's content
        for section_name, section_content in structured_data.get('report', {}).get('sections', {}).items():
            if isinstance(section_content, dict):
                for subsection_name, subsection_content in section_content.items():
                    # Add section content
                    requests.append({
                        'insertText': {
                            'location': {'index': current_index},
                            'text': subsection_content + '\n'
                        }
                    })
                    
                    # Add footnotes for this section
                    for footnote in footnotes:
                        if footnote['location'] in subsection_content:
                            requests.append({
                                'createFootnote': {
                                    'location': {
                                        'index': current_index + subsection_content.index(footnote['location']) + 
                                                len(footnote['location'])
                                    },
                                    'text': f"Source: {footnote['url']}, accessed on {footnote['access_date']}"
                                }
                            })
                    
                    current_index += len(subsection_content) + 1
            else:
                # Add section content
                requests.append({
                    'insertText': {
                        'location': {'index': current_index},
                        'text': section_content + '\n'
                    }
                })
                
                # Add footnotes for this section
                for footnote in footnotes:
                    if footnote['location'] in section_content:
                        requests.append({
                            'createFootnote': {
                                'location': {
                                    'index': current_index + section_content.index(footnote['location']) + 
                                            len(footnote['location'])
                                },
                                'text': f"Source: {footnote['url']}, accessed on {footnote['access_date']}"
                            }
                        })
                
                current_index += len(section_content) + 1

        # Execute all requests
        if requests:
            service.documents().batchUpdate(
                documentId=document_id,
                body={'requests': requests}
            ).execute()

        return f"https://docs.google.com/document/d/{document_id}/edit"

    def _download_from_google_docs(self, gdoc_url: str) -> str:
        """Download Google Doc back to Word format."""
        # Extract document ID from URL
        doc_id = gdoc_url.split('/')[-2]
        
        drive_service = build('drive', 'v3', credentials=self.google_creds)
        
        # Export as Word document
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = f"Subcontractor_Report_Final_{timestamp}.docx"
        
        request = drive_service.files().export_media(
            fileId=doc_id,
            mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        with open(output_path, 'wb') as f:
            f.write(request.execute())
        
        return output_path